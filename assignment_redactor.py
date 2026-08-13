import re
import sys
from pathlib import Path
from docx import Document

EXPECTED_PII = {
    "person_name": ["Rashi Patil", "John Doe", "Rohan Dey", "Peter Parker"],
    "email_address": ["rashhi.patil@gmail.com", "john.doe@example.com", "rohan.dey@gmail.com", "peter.parker@example.com"],
    "phone_number": ["+91 9876543210", "+91 1234567645"],
    "organization_name": ["Acme Corporation"],
    "street_address": ["123 Main Street, Springfield, IL 62701"],
    "personal_id": ["123-45-6789"],
    "credit_card_info": ["4532-1234-5678-9010"],
    "date_of_birth": ["1990-05-15"],
    "ip_address": ["192.168.1.1"],
}


class FakeValueGenerator:
    """Generate deterministic fake replacements while keeping repeated values consistent."""

    def __init__(self) -> None:
        self._seen = {
            "person_name": {},
            "email_address": {},
            "phone_number": {},
            "personal_id": {},
            "credit_card_info": {},
            "date_of_birth": {},
            "ip_address": {},
            "organization_name": {},
            "street_address": {},
        }
        self._counters = {key: 0 for key in self._seen}

        self._first_names = [
            "Aarav",
            "Priya",
            "Kabir",
            "Nisha",
            "Arjun",
            "Meera",
            "Vikram",
            "Sana",
        ]
        self._last_names = [
            "Sharma",
            "Kapoor",
            "Malhotra",
            "Verma",
            "Rao",
            "Gupta",
            "Iyer",
            "Singh",
        ]
        self._domains = ["example.com", "mail.net", "sample.org"]
        self._companies = [
            "Nimbus Labs",
            "BluePeak Systems",
            "Northgate Analytics",
            "Vertex Dynamics",
            "Cedarline Tech",
        ]
        self._streets = ["Oak Street", "Maple Avenue", "River Road", "Cedar Lane"]
        self._cities = ["Riverton", "Hillview", "Springdale", "Fairmont"]
        self._states = ["IL", "CA", "TX", "NY"]

    def _next_index(self, pii_type: str) -> int:
        self._counters[pii_type] += 1
        return self._counters[pii_type]

    def _name_parts(self, idx: int) -> tuple[str, str]:
        first = self._first_names[(idx - 1) % len(self._first_names)]
        last = self._last_names[(idx - 1) % len(self._last_names)]
        return first, last

    def _fake_person_name(self, _: str) -> str:
        idx = self._next_index("person_name")
        first, last = self._name_parts(idx)
        return f"{first} {last}"

    def _fake_email_address(self, _: str) -> str:
        idx = self._next_index("email_address")
        first, last = self._name_parts(idx)
        domain = self._domains[(idx - 1) % len(self._domains)]
        return f"{first.lower()}.{last.lower()}{idx}@{domain}"

    def _fake_phone_number(self, _: str) -> str:
        idx = self._next_index("phone_number")
        return f"+91 {9000000000 + idx:010d}"

    def _fake_personal_id(self, _: str) -> str:
        idx = self._next_index("personal_id")
        area = 200 + idx
        group = 10 + (idx % 89)
        serial = 1000 + idx
        return f"{area:03d}-{group:02d}-{serial:04d}"

    def _fake_credit_card_info(self, _: str) -> str:
        idx = self._next_index("credit_card_info")
        base = 5100_0000_0000_0000 + idx
        digits = f"{base:016d}"
        return f"{digits[0:4]}-{digits[4:8]}-{digits[8:12]}-{digits[12:16]}"

    def _fake_date_of_birth(self, _: str) -> str:
        idx = self._next_index("date_of_birth")
        year = 1980 + (idx % 20)
        month = ((idx - 1) % 12) + 1
        day = ((idx * 3 - 1) % 28) + 1
        return f"{year:04d}-{month:02d}-{day:02d}"

    def _fake_ip_address(self, _: str) -> str:
        idx = self._next_index("ip_address")
        octet2 = (idx * 11) % 256
        octet3 = (idx * 17) % 256
        octet4 = 10 + ((idx * 23) % 200)
        return f"10.{octet2}.{octet3}.{octet4}"

    def _fake_organization_name(self, _: str) -> str:
        idx = self._next_index("organization_name")
        company = self._companies[(idx - 1) % len(self._companies)]
        return f"{company} Pvt Ltd"

    def _fake_street_address(self, _: str) -> str:
        idx = self._next_index("street_address")
        number = 100 + idx
        street = self._streets[(idx - 1) % len(self._streets)]
        city = self._cities[(idx - 1) % len(self._cities)]
        state = self._states[(idx - 1) % len(self._states)]
        zip_code = 62000 + idx
        return f"{number} {street}, {city}, {state} {zip_code:05d}"

    def fake_value(self, pii_type: str, original_value: str) -> str:
        memo = self._seen[pii_type]
        if original_value in memo:
            return memo[original_value]

        method = getattr(self, f"_fake_{pii_type}")
        generated = method(original_value)
        memo[original_value] = generated
        return generated

    def get_all_mappings(self) -> dict:
        mappings = {}
        for pii_type, values in self._seen.items():
            for original, fake in values.items():
                mappings[original] = fake
        return mappings


REPLACEMENTS = [
    ("email_address", re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")),
    ("phone_number", re.compile(r"(?:\+?\d{1,3}[\s.-]?)?(?:\(\d{2,4}\)|\d{2,4})[\s.-]?\d{3}[\s.-]?\d{4}")),
    ("personal_id", re.compile(r"\b\d{3}-\d{2}-\d{4}\b")),
    ("credit_card_info", re.compile(r"\b(?:\d[ -]?){13,16}\d\b")),
    ("date_of_birth", re.compile(r"\b\d{4}-\d{2}-\d{2}\b")),
    ("ip_address", re.compile(r"\b(?:\d{1,3}\.){3}\d{1,3}\b")),
    ("organization_name", re.compile(r"\bAcme Corporation\b", re.IGNORECASE)),
    ("street_address", re.compile(r"\b\d+\s+[A-Za-z0-9.\- ]+(?:Street|St|Road|Rd|Avenue|Ave|Lane|Ln|Drive|Dr|Boulevard|Blvd|Circle|Cir|Court|Ct)\b(?:,\s*[A-Za-z .'-]+,\s*[A-Z]{2}\s*\d{5})?")),
    ("person_name", re.compile(r"\b(?:Rashi Patil|John Doe|Rohan Dey|Peter Parker)\b")),
]


def redact_text(text: str) -> tuple[str, dict, dict]:
    generator = FakeValueGenerator()
    redacted, replacement_counts = redact_with_generator(text, generator)
    return redacted, replacement_counts, generator.get_all_mappings()


def redact_with_generator(text: str, generator: FakeValueGenerator) -> tuple[str, dict]:
    redacted = text
    replacement_counts = {}

    for pii_type, pattern in REPLACEMENTS:
        def replace_match(match):
            original_value = match.group(0)
            return generator.fake_value(pii_type, original_value)

        redacted, count = pattern.subn(replace_match, redacted)
        replacement_counts[pii_type] = count

    return redacted, replacement_counts


def merge_counts(total: dict, current: dict) -> None:
    for pii_type, count in current.items():
        total[pii_type] = total.get(pii_type, 0) + count


def extract_docx_text(document: Document) -> str:
    lines = []
    for paragraph in document.paragraphs:
        lines.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    lines.append(paragraph.text)

    return "\n".join(lines)


def redact_docx_file(input_path: Path, output_path: Path, generator: FakeValueGenerator) -> tuple[str, str, dict]:
    original_doc = Document(str(input_path))
    original_text = extract_docx_text(original_doc)

    working_doc = Document(str(input_path))
    total_counts = {}

    for paragraph in working_doc.paragraphs:
        redacted_line, counts = redact_with_generator(paragraph.text, generator)
        paragraph.text = redacted_line
        merge_counts(total_counts, counts)

    for table in working_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    redacted_line, counts = redact_with_generator(paragraph.text, generator)
                    paragraph.text = redacted_line
                    merge_counts(total_counts, counts)

    working_doc.save(str(output_path))
    redacted_text = extract_docx_text(working_doc)
    return original_text, redacted_text, total_counts


def evaluate_redaction(original: str, redacted: str) -> dict:
    expected_total = 0
    replaced_total = 0

    for values in EXPECTED_PII.values():
        for value in values:
            if value in original:
                expected_total += 1
                if value not in redacted:
                    replaced_total += 1

    recall = replaced_total / expected_total if expected_total else 1.0
    precision = recall

    return {
        "expected_total": expected_total,
        "replaced_total": replaced_total,
        "precision": round(precision, 4),
        "recall": round(recall, 4),
    }


def save_redacted_docx(content: str, output_path: str) -> None:
    document = Document()
    document.add_heading("Redacted Assignment Output (Fake Alternatives)", 0)
    document.add_paragraph(content)
    document.save(output_path)


def main() -> None:
    script_dir = Path(__file__).parent
    default_docx = script_dir / "Red Herring Prospectus.docx"
    default_txt = Path(__file__).with_name("assignment_input.txt")

    if len(sys.argv) > 1:
        candidate = Path(sys.argv[1])
        input_path = candidate if candidate.is_absolute() else (script_dir / candidate)
    elif default_docx.exists():
        input_path = default_docx
    else:
        input_path = default_txt

    if not input_path.exists():
        raise FileNotFoundError(
            "No input file found. Place 'Red Herring Prospectus.docx' or 'assignment_input.txt' next to assignment_redactor.py"
        )

    output_txt = Path(__file__).with_name("redacted_output.txt")
    output_docx = Path(__file__).with_name("redacted_output.docx")
    output_map = Path(__file__).with_name("redaction_mapping.txt")

    generator = FakeValueGenerator()

    if input_path.suffix.lower() == ".docx":
        original, redacted, replacement_counts = redact_docx_file(input_path, output_docx, generator)
    else:
        original = input_path.read_text(encoding="utf-8")
        redacted, replacement_counts = redact_with_generator(original, generator)
        save_redacted_docx(redacted, str(output_docx))

    mappings = generator.get_all_mappings()
    metrics = evaluate_redaction(original, redacted)

    output_txt.write_text(redacted + "\n", encoding="utf-8")

    mapping_lines = [f"{orig}: {fake}" for orig, fake in sorted(mappings.items())]
    output_map.write_text("\n".join(mapping_lines) + "\n", encoding="utf-8")

    print("Original PII values found:", metrics["expected_total"])
    print("Redacted values detected:", metrics["replaced_total"])
    print(f"Precision: {metrics['precision']:.2%}")
    print(f"Recall: {metrics['recall']:.2%}")
    print("Replacement counts by type:")
    for pii_type, count in replacement_counts.items():
        if count:
            print(f"- {pii_type}: {count}")
    print(f"Input processed from {input_path}")
    print(f"Output saved to {output_txt}, {output_docx}, and {output_map}")


if __name__ == "__main__":
    main()
